import os
import subprocess
from docx import Document
import csv

# Load the template document once
def load_template_document(form_path):
    return Document(form_path)

def fill_financial_form(document, data):
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(value))

def process_row(i, row, form_path, output_dir):
    temp_docx_path = os.path.join('output_docx', f'filled_form_{i+1}.docx')
    filled_document = load_template_document(form_path)
    fill_financial_form(filled_document, row)
    filled_document.save(temp_docx_path)
    output_pdf = os.path.join(output_dir, os.path.basename(temp_docx_path).replace('.docx', '.pdf'))
    try:
        subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, temp_docx_path], check=True)
        print(f"Conversion successful. PDF saved to: {output_pdf}")
    except subprocess.CalledProcessError as e:
        print(f"Conversion failed for row {i+1}: {e}")
        # Remove the generated DOCX file if conversion fails

def fill_forms_from_csv(csv_path, form_path):
    with open(csv_path, 'r') as csv_file:
        csv_reader = csv.DictReader(csv_file)
        output_dir = 'output_pdfs'
        os.makedirs(output_dir, exist_ok=True)
        for i, row in enumerate(csv_reader):
            # if i == 10:
            #     return 
            process_row(i, row, form_path, output_dir)

# Path to the empty financial form and CSV file
form_path = 'test_template.docx'
csv_path = 'Database/financial_data.csv'

# Fill the forms with data from CSV and save as PDF
fill_forms_from_csv(csv_path, form_path)
print("Scucessfully Generated..........")
